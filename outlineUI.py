#*************************************************************************
# Develop of Courses Outline 
# Auckland Institute of Studies
# Developers: William Martin, June , Sun....
# Date of creation 03/11/2021
#*************************************************************************
from tkinter.filedialog import askopenfile,askdirectory,asksaveasfile
from tkinter import *
from tkinter.ttk import Progressbar,Button
from PIL import ImageTk, Image
from tkinter import ttk , messagebox
import tkinter as tk,time, re 
import docx,shutil,os,globalVars as gv
import pandas as pd
from docx.shared import RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Root(Tk):
    def __init__(self):
        super(Root, self).__init__()
        gv.vcmd = (self.register(self.isNumeric))

        #Window mesures:
        self.title("Course Outline Generator")
        self.geometry('700x500')
        self.minsize(640,600)
        self.resizable(False, False)
        #Form limits:
        self.rbCoursePerson=tk.StringVar()
        self.rbCoursePerson.set('Lecturer')
        self.nameValue = StringVar()
        self.nameValue.trace("w", lambda *args: self.character_limit(self.nameValue, gv.mxNa))
        self.nameLValue = StringVar()
        self.nameLValue.trace("w", lambda *args: self.character_limit(self.nameLValue, gv.mxNa))
        self.ph1Value = StringVar()
        self.ph1Value.trace("w", lambda *args: self.field_limit(self.ph1Value, gv.mxph1))
        self.ph3Value = StringVar()
        self.ph3Value.trace("w", lambda *args: self.field_limit(self.ph3Value, gv.mxph3))
        self.emailValue = StringVar()
        self.emailValue.trace("w", lambda *args: self.characterNum_limit(self.emailValue, gv.mxEm))
        self.roomValue = StringVar()
        self.roomValue.trace("w", lambda *args: self.alphaNum_limit(self.roomValue, gv.mxRo))
        self.hourValue = StringVar()
        self.hourValue.trace("w", lambda *args: self.check_hour(self.hourValue, 2))
        self.minValue = StringVar()
        self.minValue.trace("w", lambda *args: self.check_minute(self.minValue, 2))
        self.hourValueF = StringVar()
        self.hourValueF.trace("w", lambda *args: self.check_hour(self.hourValueF, 2))
        self.minValueF = StringVar()
        self.minValueF.trace("w", lambda *args: self.check_minute(self.minValueF, 2))
        self.empty=tk.StringVar()
#<editor-fold desc="Logic Functions">
    def uploadFiles(self):
        if gv.fName != "":
            gv.lblFile.config(text=gv.fName)
            pb1 = Progressbar(
                    root,
                    orient=HORIZONTAL,
                    length=300,
                    mode='determinate'
                    )
            pb1.place(x=gv.xPosF+120,y=gv.yBtnPos)
            labelPorc = ttk.Label(root,font=("bold",8))
            labelPorc.place(x=gv.xPosF+260,y=gv.yBtnPos+2)
            for i in range(6):
                root.update_idletasks()
                pb1['value'] += 20
                labelPorc.config(text=str(int(pb1['value'])-20) + "%")
                time.sleep(0.1)
            labelPorc.destroy()
            pb1.destroy()
            gv.target = 'docs'
            try:
                shutil.copy(gv.original, gv.target)
                print("File copied successfully.")
                gv.entry_name.focus_set()
                messagebox.showinfo(title="INFORMATION", message='File Uploaded Successfully!')
            # If source and destination are same
            except shutil.SameFileError:
                print("Source and destination represents the same file.")
                messagebox.showinfo(title="Error:", message='Source and destination represents the same file.')

            # If there is any permission issue
            except PermissionError:
                print("Permission denied.")
                messagebox.showinfo(title="Error:", message='Permission denied.')

            # For other errors
            except Exception as ep:
                print("Error occurred while copying file.")
                messagebox.showerror(title="Error:", message='Error occurred while copying file.')
        else:
            messagebox.showerror(title="INFORMATION", message='Insert the File!')

    def open_file(self):
        file_path = askopenfile(mode='r', filetypes=[('Doc Files', '*docx')])
        if file_path is not None:
            gv.original=file_path.name
            pos = gv.original.rfind("/")
            gv.fName=(gv.original[pos+1:len(gv.original)+1])
            gv.fileSize = os.path.getsize(gv.original)
            print('Size del archivo: ' + str(gv.fileSize))
            if(int(gv.fileSize)<2097153):
                gv.state = True
                self.ableControls(gv.state)
                self.uploadFiles()
            else: messagebox.showwarning(title="INFORMATION", message='File size too big, choose a file with size < 2 Mb!')
        else:
            if gv.original == "":
                gv.state = False
                self.ableControls(gv.state)
                messagebox.showerror(title="INFORMATION", message='No file chosen, Insert the file again!')

    def segmentLine(self):
        w = tk.Canvas(self, width=690, height=3)
        w.place(y=gv.yPos,x=0)
        w.config( background=gv.bgSL)

    def update_click(self):
        self.upd_Docum_Docm()
        self.upd_Docum_Frm()

    def set_Final_Formval(self):

        gv.nameF = gv.entry_name.get().title() + " "+ gv.entry_Lname.get().title()
        gv.roomF = gv.entry_room.get().upper()
        gv.phoneF = gv.entry_phone1.get() + " ext. " +  gv.entry_ext.get()
        gv.hourI= gv.hourI="0"+gv.entry_contHour.get()  if len(gv.entry_contHour.get())==1 else gv.entry_contHour.get()
        gv.hourI= gv.hourI+":0"+gv.entry_contMinute.get() if len(gv.entry_contMinute.get())==1 else gv.hourI+":"+gv.entry_contMinute.get()
        gv.hourF= gv.hourF="0"+gv.entry_contHourF.get()  if len(gv.entry_contHourF.get())==1 else gv.entry_contHourF.get()
        gv.hourF= gv.hourF+":0"+gv.entry_contMinuteF.get() if len(gv.entry_contMinuteF.get())==1 else gv.hourF+":"+gv.entry_contMinuteF.get()

        if(gv.hourI>=gv.hourF):
            messagebox.showinfo("Bad Time Format","The End time must be greater than Start time, please correct it.")
            gv.entry_contHourF.focus_set()
            return False

        gv.hourI = gv.hourI+" AM" if gv.hourI[:1]=="0" else gv.hourI
        gv.hourI = gv.hourI+" PM" if gv.hourI[:2]=="12" else gv.hourI
        gv.hourF = gv.hourF+" AM" if gv.hourF[:1]=="0" else gv.hourF
        gv.hourF = gv.hourF+" PM" if gv.hourF[:2]=="12" else gv.hourF
        gv.contactHour = gv.hourI + " to " + gv.hourF

        print("Person: " + self.rbCoursePerson.get())
        print("Trimestre: %s\nYear: %s" % (gv.trimester_cb.get(), gv.year_cb.get()))
        print("Name: %s\nEmail: %s" % (gv.nameF, gv.entry_email.get()))
        print("Room: %s\nPhone: %s" % (gv.roomF, gv.phoneF))
        print("Hour I: " + gv.hourI)
        print("Hour F: " + gv.hourF)
        print("Contact Hour: " + gv.contactHour)
        return True

    def upd_Docum_Frm(self):
        res = False
        if(self.validateForm()): res = self.set_Final_Formval()
        valuesDocx = [gv.nameF, gv.roomF,gv.phoneF, gv.entry_email.get(),gv.contactHour]
        flag=0
        fieldsCopied=0
        try:
            if res and len(gv.fName)>0:
                #document = docx.Document(gv.originalDoc)
                #Jess changed here 07/11
                document = gv.CO_Doc
                for par in document.paragraphs:  # to extract the whole text
                    if self.rbCoursePerson.get()+":" in par.text:
                        flag=1
                    if flag==1:
                        for i in range(len(gv.fieldsDocx)):
                            if gv.fieldsDocx[i] in par.text:
                                fieldsCopied+=1
                                if len(valuesDocx[i])>0:
                                    tmp_text = par.text
                                    if "ext.#" in gv.fieldsDocx[i]:
                                        if gv.entry_phone1.get()!="":
                                            tmp_text = tmp_text.replace(gv.fieldsDocx[i],valuesDocx[i])
                                            par.text=tmp_text
                                            break
                                    else:
                                        tmp_text = tmp_text.replace(gv.fieldsDocx[i],valuesDocx[i])
                                        par.text=tmp_text
                                        break
                        if fieldsCopied==len(gv.fieldsDocx) and flag==1: break
                document.save('docs/'+gv.targetDoc)
                fname = gv.entry_name.get().title()
                lname = gv.entry_Lname.get().title()
                file_name = gv.courseCode + "-" + gv.trimester_cb.get() + "-" + gv.year_cb.get() +"-CourseOutline-draft1-"+fname[0]+lname[0]
                file_path = asksaveasfile(mode='w', filetypes=[('Doc Files', '*docx')], initialfile = file_name+".docx")
                if file_path is not None:
                    gv.state = True
                    gv.CO_Doc.save(file_path.name)
                    if self.empty.get()=="1":
                        self.clearControls()
                    # messagebox.showinfo(title="Successful process", message='File updated, Course Outline generated.')
                    if messagebox.askokcancel("Successful process","File updated, Course Outline generated.\nDo you want to open this file?"):
                        print("Escogio: "+str(res))
                        os.startfile(gv.originalDoc)

        except docx.opc.exceptions.PackageNotFoundError:
            messagebox.showerror(title="Error:", message='The document docOrigin/TempleteCO.docx is not accesible. Please verify that it is in the folder.')
        except PermissionError:
            messagebox.showerror(title="Error:", message='The document is not accesible. It can be open if so close it, please check.')

#</editor-fold>

#<editor-fold desc="Constructors">
    def createHeader(self):
        heading = tk.Label(self, text="Course Outline Generator")
        heading.config(font=(gv.tFont, gv.mtSize),fg=gv.stColor)
        heading.pack(padx=50, pady=65)
        uplHeading = tk.Label(root, text="Course Descriptor")
        uplHeading.config(font=(gv.tFont, gv.stSize),fg=gv.stColor)
        uplHeading.place(x=gv.xPosL, y=gv.yPos)
        gv.yPos+=30
        gv.upld = tk.Button(self,text='Upload a Course Descriptor' , width=30,bg="blue",fg='white',activebackground='#0052cc', activeforeground='#aaffaa',command = self.open_file )
        gv.upld.place(x=gv.xPosL,y=gv.yPos)
        gv.yBtnPos = gv.yPos
        #File label:
        gv.lblFile = tk.Label(self,textvariable=gv.fName, width=35,font=(gv.lbFont,gv.lbSize),fg='blue')
        gv.lblFile.place(x=gv.xPosL+240,y=gv.yPos+2)
        gv.yPos += 30
        label_upl = tk.Label(self,text="Accepted file types: .doc and .docx (2MB limit)", width=35,font=(gv.lbFont,gv.lbSize))
        label_upl.place(x=gv.xPosL,y=gv.yPos)
        gv.yPos += 30
        self.segmentLine()
        gv.yPos += 20

    def createPeriod(self):
        label_trimester =tk.Label(self,text="Trimester", width=8,font=(gv.lbFont,gv.lbSize))
        label_trimester.place(x=gv.xCbxLbl,y=gv.yPos)
        gv.xCbxLbl+=100
        gv.trimester_cb = ttk.Combobox(self)
        gv.trimester_cb['values'] = gv.months
        gv.trimester_cb['state'] = 'readonly'
        gv.trimester_cb.current(1)
        gv.trimester_cb.config(width=10)
        gv.trimester_cb.place(x=gv.xCbxLbl,y=gv.yPos)
        gv.xCbxLbl+=110
        label_year =tk.Label(root,text="Year", width=9,font=(gv.lbFont,gv.lbSize))
        label_year.place(x=gv.xCbxLbl,y=gv.yPos) 
        gv.xCbxLbl+=100
        gv.year_cb = ttk.Combobox(self)
        gv.year_cb['values'] = gv.years
        gv.year_cb['state'] = 'readonly'
        gv.year_cb.config(width=10)
        gv.year_cb.current(1)
        gv.year_cb.place(x=gv.xCbxLbl,y=gv.yPos)
        gv.yPos+=33
        self.segmentLine()
        gv.yPos+=20

    def createLecturerInf(self):
        lectHeading = tk.Label(self, text="Lecturer Information")
        lectHeading.config(font=(gv.tFont, gv.stSize),fg=gv.stColor)
        lectHeading.place(x=gv.xPosL, y=gv.yPos)
        gv.yPos+=30
        self.rb1=tk.Radiobutton(self,text="Lecturer",padx= 5, variable=self.rbCoursePerson, value="Lecturer").place(x=gv.xPosL,y=gv.yPos)
        self.rb2=tk.Radiobutton(self,text="Course Coordinator",padx= 20, variable=self.rbCoursePerson, value="Course Coordinator").place(x=gv.xPosF,y=gv.yPos)
        gv.yPos+=40
        label_name =tk.Label(self,text="First & Last Name", width=14,font=(gv.lbFont,gv.lbSize))
        label_name.place(x=gv.xPosL,y=gv.yPos)
        gv.entry_name=tk.Entry(self , textvariable = self.nameValue)
        gv.entry_name.config(width=18)
        gv.entry_name.place(x=gv.xPosF,y=gv.yPos)
        gv.entry_Lname=tk.Entry(self , textvariable = self.nameLValue)
        gv.entry_Lname.config(width=18)
        gv.entry_Lname.place(x=gv.xPosF+130,y=gv.yPos)
        label_nameM =tk.Label(self,text="Max. ("+str(gv.mxNa)+" char.)", fg=gv.lbCColor, width=12,font=(gv.lbFont,gv.lbSize-2))
        label_nameM.place(x=gv.xPosF+265,y=gv.yPos+1)
        gv.yPos+=30
        label_room =tk.Label(self,text="Room", width=5,font=(gv.lbFont,gv.lbSize))
        label_room.place(x=gv.xPosL,y=gv.yPos)
        gv.entry_room=tk.Entry(self,textvariable = self.roomValue)
        gv.entry_room.config(width=40)
        gv.entry_room.place(x=gv.xPosF,y=gv.yPos)
        label_roomM =tk.Label(self,text="Max. ("+str(gv.mxRo)+" char.)", fg=gv.lbCColor, width=12,font=(gv.lbFont,gv.lbSize-2))
        label_roomM.place(x=gv.xPosF+265,y=gv.yPos+1)
        gv.yPos+=30
        label_Phone = tk.Label(self,text="Phone", width=5,font=(gv.lbFont,gv.lbSize))
        label_Phone.place(x=gv.xPosL,y=gv.yPos)
        label_PhoneInd = tk.Label(self,text="(+64)", width=5,font=(gv.lbFont,gv.lbSize-1))
        label_PhoneInd.place(x=gv.xPosL+79,y=gv.yPos)
        gv.entry_phone1=tk.Entry(self, textvariable = self.ph1Value, validate='all',validatecommand=(gv.vcmd, '%P'))
        gv.entry_phone1.config(width=10)
        gv.entry_phone1.place(x=gv.xPosF,y=gv.yPos)
        label_dashExt =tk.Label(self,text=" ext", width=3,font=(gv.lbFont,gv.lbSize))
        label_dashExt.place(x=gv.xPosF+63,y=gv.yPos-2)
        gv.entry_ext=tk.Entry(self, textvariable = self.ph3Value, validate='all',validatecommand=(gv.vcmd, '%P'))
        gv.entry_ext.config(width=5)
        gv.entry_ext.place(x=gv.xPosF+95,y=gv.yPos)
        label_phoneEx =tk.Label(self,text="p.e. (2013451717 ext 1245)", fg=gv.lbCColor, width=21,font=(gv.lbFont,gv.lbSize-2))
        label_phoneEx.place(x=gv.xPosF+268,y=gv.yPos+1)
        gv.yPos+=30
        label_email =tk.Label(self,text="Email", width=5,font=(gv.lbFont,gv.lbSize))
        label_email.place(x=gv.xPosL,y=gv.yPos)
        gv.entry_email=tk.Entry(self, textvariable = self.emailValue)
        gv.entry_email.config(width=40)
        gv.entry_email.place(x=gv.xPosF,y=gv.yPos)
        label_emailM =tk.Label(self,text="Max. ("+str(gv.mxEm)+" char.)", fg=gv.lbCColor, width=12,font=(gv.lbFont,gv.lbSize-2))
        label_emailM.place(x=gv.xPosF+265,y=gv.yPos+1)
        gv.yPos+=30
        label_ContactH = tk.Label(self,text="Contact hour", width=10,font=(gv.lbFont,gv.lbSize))
        label_ContactH.place(x=gv.xPosL,y=gv.yPos)
        gv.entry_contHour=tk.Entry(self, textvariable = self.hourValue, validate='all',validatecommand=(gv.vcmd, '%P'))
        gv.entry_contHour.config(width=3)
        gv.entry_contHour.place(x=gv.xPosF,y=gv.yPos)
        label_dash =tk.Label(self,text=":", width=1,font=(gv.lbFont,gv.lbSize))
        label_dash.place(x=gv.xPosF+20,y=gv.yPos-2)
        gv.entry_contMinute=tk.Entry(self, textvariable = self.minValue, validate='all',validatecommand=(gv.vcmd, '%P'))
        gv.entry_contMinute.config(width=3)
        gv.entry_contMinute.place(x=gv.xPosF+34,y=gv.yPos)
        label_dashhf =tk.Label(self,text="To", width=1,font=(gv.lbFont,gv.lbSize))
        label_dashhf.place(x=gv.xPosF+64,y=gv.yPos-2)
        gv.entry_contHourF=tk.Entry(self, textvariable = self.hourValueF, validate='all',validatecommand=(gv.vcmd, '%P'))
        gv.entry_contHourF.config(width=3)
        gv.entry_contHourF.place(x=gv.xPosF+84,y=gv.yPos)
        label_dashmf =tk.Label(self,text=":", width=1,font=(gv.lbFont,gv.lbSize))
        label_dashmf.place(x=gv.xPosF+104,y=gv.yPos-2)
        gv.entry_contMinuteF=tk.Entry(self, textvariable = self.minValueF, validate='all',validatecommand=(gv.vcmd, '%P'))
        gv.entry_contMinuteF.config(width=3)
        gv.entry_contMinuteF.place(x=gv.xPosF+118,y=gv.yPos)
        label_year =tk.Label(root,text="Day", width=3,font=(gv.lbFont,gv.lbSize))
        label_year.place(x=gv.xPosF+147,y=gv.yPos-2)
        gv.day_cb = ttk.Combobox(self)
        gv.day_cb['values'] = gv.days
        gv.day_cb['state'] = 'readonly'
        gv.day_cb.config(width=10)
        gv.day_cb.current(1)
        gv.day_cb.place(x=gv.xPosF+179,y=gv.yPos-1)

        label_hour =tk.Label(self,text="MM:HH p.e. 09:30 (AM) To 15:45 (PM)", fg=gv.lbCColor, width=30,font=(gv.lbFont,gv.lbSize-2))
        label_hour.place(x=gv.xPosF+270,y=gv.yPos)
        gv.yPos+=35
        self.segmentLine()
        gv.yPos+=30

    def endForm(self):
        #gv.state = False
        self.ableControls(gv.state)
        tk.Button(self, text='Download a Course Outline' , width=30,bg="green",fg='white',activebackground='#0052cc', activeforeground='#aaffaa', command=self.update_click).place(x=gv.xPosF+90,y=gv.yPos)
        gv.yPos+=30
        gv.empty_ch=tk.Checkbutton(self, text='Empty the Lecturer Information fields after download',variable=self.empty, onvalue="1", offvalue="0" )
        gv.empty_ch.select()
        gv.empty_ch.place(x=gv.xPosF+40,y=gv.yPos)
        gv.upld.focus_set()

#</editor-fold>

#<editor-fold desc="Form Functions">
    def ableControls(self,state):
        if state:
            gv.trimester_cb['state'] = tk.NORMAL
            gv.year_cb['state'] = tk.NORMAL
            gv.entry_name['state'] = tk.NORMAL
            gv.entry_Lname['state'] = tk.NORMAL
            gv.entry_room['state'] = tk.NORMAL
            gv.entry_phone1['state'] = tk.NORMAL
            gv.entry_ext['state'] = tk.NORMAL
            gv.entry_email['state'] = tk.NORMAL
            gv.entry_contHour['state'] = tk.NORMAL
            gv.entry_contMinute['state'] = tk.NORMAL
            gv.entry_contHourF['state'] = tk.NORMAL
            gv.entry_contMinuteF['state'] = tk.NORMAL
            gv.day_cb['state'] = tk.NORMAL
        else:
            gv.trimester_cb['state'] = tk.DISABLED
            gv.year_cb['state'] = tk.DISABLED
            gv.entry_name['state'] = tk.DISABLED
            gv.entry_Lname['state'] = tk.DISABLED
            gv.entry_room['state'] = tk.DISABLED
            gv.entry_phone1['state'] = tk.DISABLED
            gv.entry_ext['state'] = tk.DISABLED
            gv.entry_email['state'] = tk.DISABLED
            gv.entry_contHour['state'] =  tk.DISABLED
            gv.entry_contMinute['state'] =  tk.DISABLED
            gv.entry_contHourF['state'] = tk.DISABLED
            gv.entry_contMinuteF['state'] = tk.DISABLED
            gv.day_cb['state'] = tk.DISABLED

    def validateForm(self):
        try:
            style = ttk.Style()
            style.configure("TCombobox", fieldbackground="yellow")
            res = False
            if gv.fName == "":
                messagebox.showinfo(title="INFORMATION", message='The File was not chosen, please choose the Course Descriptor!')
                self.ableControls(False)
                gv.trimester_cb.focus_set()
                return False
            if gv.trimester_cb.get()=='Select >':
                messagebox.showinfo(title="INFORMATION", message='The Trimester was not chosen, choose the Trimester!')
                gv.trimester_cb.focus_set()
                return False
            if gv.year_cb.get()=='Select >':
                messagebox.showinfo(title="INFORMATION", message='The Year was not chosen, choose the Year!')
                gv.year_cb.focus_set()
                return False
            if len(gv.entry_name.get().strip()) == 0:
                messagebox.showinfo(title="INFORMATION", message='The Name is empty, please enter It!')
                gv.entry_name.focus_set()
                return False
            elif (self.validateName(gv.entry_name.get().strip().title(),gv.nameRegex)==False):
                messagebox.showinfo(title="INFORMATION", message='The Name structure is bad formed, enter it again!')
                gv.entry_name.focus_set()
                return False
            if len(gv.entry_Lname.get().strip()) == 0:
                messagebox.showinfo(title="INFORMATION", message='The Last Name is empty, please enter It!')
                gv.entry_Lname.focus_set()
                return False
            elif (self.validateName(gv.entry_Lname.get().strip().title(),gv.nameRegex)==False):
                messagebox.showinfo(title="INFORMATION", message='The Last Name structure is bad formed, enter it again!')
                gv.entry_Lname.focus_set()
                return False
            if len(gv.entry_room.get().strip()) == 0:
                messagebox.showinfo(title="INFORMATION", message='The Room is empty, please enter It!')
                gv.entry_room.focus_set()
                return False
            if len(gv.entry_phone1.get().strip()) == 0:
                messagebox.showinfo(title="INFORMATION", message='The Phone is empty, please enter It!')
                gv.entry_phone1.focus_set()
                return False
            if len(gv.entry_email.get().strip()) == 0:
                messagebox.showinfo(title="INFORMATION", message='The Email is empty, please enter It!')
                gv.entry_email.focus_set()
                return False
            if self.check(gv.entry_email.get()) == False:
                messagebox.showinfo(title="INFORMATION", message='The Email is badly formed, please correct It!')
                gv.entry_email.focus_set()
                return False
            if len(gv.entry_contHour.get().strip()) == 0:
                messagebox.showinfo(title="INFORMATION", message='The initial Hour is incomplete, please enter It!')
                gv.entry_contHour.focus_set()
                return False
            if len(gv.entry_contMinute.get().strip()) == 0:
                messagebox.showinfo(title="INFORMATION", message='The initial Minute is incomplete, please enter It!')
                gv.entry_contMinute.focus_set()
                return False
            if len(gv.entry_contHourF.get().strip()) == 0:
                messagebox.showinfo(title="INFORMATION", message='The final Hour is incomplete, please enter It!')
                gv.entry_contHourF.focus_set()
                return False
            if len(gv.entry_contMinuteF.get().strip()) == 0:
                messagebox.showinfo(title="INFORMATION", message='The final Minute is incomplete, please enter It!')
                gv.entry_contMinuteF.focus_set()
                return False
            if gv.day_cb.get()=='Select >':
                messagebox.showinfo(title="INFORMATION", message='The Day was not chosen, choose the day of contact!')
                gv.day_cb.focus_set()
                return False
            else:
                return True
        except Exception as ep:
            messagebox.showerror("Error:", ep)
        return res

    def clearControls(self):
        gv.entry_name.delete(0,"end")
        gv.entry_name.insert(0, '') 
        gv.entry_Lname.delete(0,"end")
        gv.entry_Lname.insert(0, '')
        gv.entry_room.delete(0,"end")
        gv.entry_room.insert(0, '')
        gv.entry_phone1.delete(0,"end")
        gv.entry_phone1.insert(0, '')
        gv.entry_ext.delete(0,"end")
        gv.entry_ext.insert(0, '')
        gv.entry_email.delete(0,"end")
        gv.entry_email.insert(0, '')
        gv.entry_contHour.delete(0,"end")
        gv.entry_contHour.insert(0, '')
        gv.entry_contMinute.delete(0,"end")
        gv.entry_contMinute.insert(0, '')
        gv.entry_contHourF.delete(0,"end")
        gv.entry_contHourF.insert(0, '')
        gv.entry_contMinuteF.delete(0,"end")
        gv.entry_contMinuteF.insert(0, '')
        gv.trimester_cb.current(1)
        gv.year_cb.current(1)
        gv.day_cb.current(1)

    def isNumeric(self, P):
        if str.isdigit(P) or P == "":
            return True
        else:
            return False

    def validateName(self,name,P):
        if(re.fullmatch(P, name)):
            return True
        else:
            return False

    def check(self, email):
        if(re.fullmatch(gv.emailRegex, email)):
            return True
        else:
            return False

    def character_limit(self, value, limit):
        value.set(value.get()[:limit])
        if re.match(gv.textRegex,value.get()) is None:
            value.set(value.get()[:-1])
        return True

    def characterNum_limit(self, value, limit):
        value.set(value.get()[:limit])
        if re.match(gv.textNumRegex,value.get()) is None:
            value.set(value.get()[:-1])
        return True

    def alphaNum_limit(self, value, limit):
        value.set(value.get()[:limit])
        if re.match(gv.alphanumRegex,value.get()) is None:
            value.set(value.get()[:-1])
        return True

    def check_hour(self, value, limit):
        value.set(value.get()[:limit])
        if re.match(gv.hourRegex,value.get()) is None:
            if len(value.get())==2:
               value.set(value.get()[:-1])
        return True

    def check_minute(self, value, limit):
        value.set(value.get()[:limit])
        if re.match(gv.minRegex,value.get()) is None:
            if len(value.get())==2:
               value.set(value.get()[:-1])
        return True

    def field_limit(self, value, limit):
        value.set(value.get()[:limit])
        return True

    def funcEnter(self,event):
        if self.focus_get() == gv.upld:
            self.open_file()
#</editor-fold>

#<editor-fold desc="Document Functions">
    
    def upd_Docum_Docm(self): 
        
        gv.CO_Doc = docx.Document(gv.originalDoc)

        inputDoc = docx.Document(gv.original)
        print(gv.original)

        def read_docx_table(doc, table_num=1, nhader=1):
            table = doc.tables[table_num-1]
            data = [[cell.text for cell in row.cells] for row in table.rows]
        
            df = pd.DataFrame(data)

            if nhader == 1 :
                df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)

            return df

        df = read_docx_table(inputDoc,1,0)
    
        #Get the first Column
        gv.firstColumn = pd.Series(df[:][0], name="s")
        #Programme
        gv.programme = df[gv.firstColumn.isin(['Programme']) == True].iloc[0, 1]
        #Course Code
        gv.courseCode = df[gv.firstColumn.isin(['Course Code']) == True].iloc[0, 1]
        #Course Title
        gv.courseTitle = df[gv.firstColumn.isin(['Course Title']) == True].iloc[0, 1]
        #NZQF Level
        gv.nzqfLevel = df[gv.firstColumn.isin(['NZQF Level']) == True].iloc[0, 1]
        #Credits
        gv.credits = df[gv.firstColumn.isin(['Credits']) == True].iloc[0, 1]
        #Prerequisites
        gv.prerequisites = df[gv.firstColumn.isin(['Prerequisites']) == True].iloc[0, 1]
        # #Co-requisites
        gv.corequisites = df[gv.firstColumn.isin(['Co-requisites']) == True].iloc[0, 1]
        # #Restrictions
        gv.restrictions = df[gv.firstColumn.isin(['Restrictions']) == True].iloc[0, 1]
        #Course Aims
        gv.courseAims = df[gv.firstColumn.isin(['Course Aims']) == True].iloc[0, 1]
        #Get rows out of 'lo' series starts its second row because the first low is 'The learners will be able to:'
        gv.learningOutcomes = df[gv.firstColumn.str.contains('Learning\nOutcomes')].iloc[1:, 2]
        #Average Learning
        gv.avgLearning = df[gv.firstColumn.str.contains('Average')].iloc[:, [1,3,4,5]]
        #Summative Assessment 
        gv.sumAssessment = df[gv.firstColumn.str.contains('Summative')].iloc[1:, [1,4,5]]  

        root.changeHeader()
        root.copySumAssesment()
        root.replaceInfo()
        
    def changeHeader(self):
        section = gv.CO_Doc.sections[0]
        header = section.header
        headerTitle = header.paragraphs[0]
        headerTitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
        headerTitle.text = gv.courseCode + "\t\tTrimester " + gv.trimester_cb.get() + ", " + gv.year_cb.get()
        headerStyle = gv.CO_Doc.styles['Header']
        headerFont = headerStyle.font
        headerFont.name = gv.tFont
        headerFont.size = Pt(12)
        headerFont.color.rgb = RGBColor(0,0,0)
        headerTitle.style = headerStyle

    def copySumAssesment(self) : 
        # for row in range(sumAssessment.shape[0]):
        #     for cell in range(sumAssessment.shape[1]):
        #         print(sumAssessment.iloc[row,cell])

        SummativeAssTbl = gv.CO_Doc.tables[3]
        
        for row in range(len(SummativeAssTbl.rows)):
            if row >= 1 and row <=4:
                for cell in range(len(SummativeAssTbl.rows[row].cells)):
                    if gv.sumAssessment.shape[0] >= row:
                        if cell == 0:
                            SummativeAssTbl.rows[row].cells[cell].text = gv.sumAssessment.iloc[row-1, 0]
                        elif cell == 1:
                            SummativeAssTbl.rows[row].cells[cell].text = gv.sumAssessment.iloc[row-1, 1]
                        elif cell == 4:
                            SummativeAssTbl.rows[row].cells[cell].text = gv.sumAssessment.iloc[row-1, 2]

    def replaceInfo(self): 
        values = gv.CO_Doc.paragraphs
        #values = docx.Document(gv.originalDoc).paragraphs
        next_ = afterNext = None
        isCourseDuration = False
        isDeleted = False
        length = len(values) 
        valuesFromCD = [gv.courseCode, gv.courseTitle, gv.prerequisites, gv.corequisites, gv.restrictions, gv.nzqfLevel, gv.credits, gv.courseAims, gv.learningOutcomes]
        
        def setStyle(style_type, targetStyle) : 
            font = targetStyle.font
            font.name = gv.tFont
            if style_type == "Heading 1":
                targetStyle.base_style = styles['Heading 1']
                font.size = Pt(18)
                font.color.rgb = RGBColor(255,255,255)
            
            elif style_type == "Normal_Bold":
                targetStyle.base_style = styles['Normal']
                font.size = Pt(11)
                font.bold = True
                font.color.rgb = RGBColor(0,0,0)

            elif style_type == "Normal":
                targetStyle.base_style = styles['Normal']
                font.size = Pt(11)
                font.color.rgb = RGBColor(0,0,0)
            
            font.italic = False
            return targetStyle 

        for index, par in enumerate(values): # to extract the whole text
            for i in range(len(gv.fieldsTitles)):
                if gv.fieldsTitles[i] in par.text:
                    styles = gv.CO_Doc.styles
                    if par.text == "+COURSE Code+" :
                        heading_style = styles.add_style('New Heading'+str(index), WD_STYLE_TYPE.PARAGRAPH)    
                        par.style = setStyle("Heading 1", heading_style)
                    elif par.text == "+COURSE Title+" :
                        heading_style = styles.add_style('New Heading'+str(index), WD_STYLE_TYPE.PARAGRAPH)
                        par.style = setStyle("Heading 1", heading_style)
                    elif "PREREQUISITES" in par.text or "CO-REQUISITES" in par.text or "RESTRICTIONS" in par.text:
                            if index < (length - 2):
                                next_ = values[index + 1]
                                afterNext = values[index + 2]
                            normal = styles.add_style('normal'+ str(index), WD_STYLE_TYPE.PARAGRAPH)
                            par.style = setStyle("Normal_Bold", normal)
                    elif "The aim of the course is to:" in par.text: 
                            if index < (length - 1):
                                next_ = values[index + 1]
                                normal = styles['Normal'] #Make Entire Normal Style same
                                par.style = next_.style = setStyle("Normal", normal)
                    elif "NZQF" in par.text or "Credits":
                        par.style = setStyle("Normal", normal) 

                    #Replace information
                    if len(valuesFromCD[i])>0:
                        tmp_text = par.text

                        if "PREREQUISITES" in tmp_text or "CO-REQUISITES" in tmp_text or "RESTRICTIONS" in tmp_text:
                            #eliminate + + and get rid of next two lines and replace it as valueDocx[i]
                            tmp_text = gv.fieldsTitles[i].replace("+","").strip()
                            next_.text = next_.text.replace(next_.text, valuesFromCD[i])
                            afterNext.text = afterNext.text.replace(afterNext.text, "")

                        if "NZQF" in tmp_text or "Credits" in tmp_text: 
                            result = tmp_text.split(':')
                            value = result[0] + ": " + valuesFromCD[i]
                            tmp_text = tmp_text.replace(gv.fieldsTitles[i], value) 
                            par.text=tmp_text
                        if "The aim of" in tmp_text :
                            next_.text = next_.text.replace(next_.text, valuesFromCD[i])
                            break
                        if "The learners will be able to:" in tmp_text :
                            autoNum = 1
                            for lo in gv.learningOutcomes:
                                
                                values[index+1].insert_paragraph_before(str(autoNum) + ".  " + lo)
                                delete_paragraph(values[index+1])
                                index = index + 1
                                autoNum = autoNum + 1
                        
                        else: 
                            tmp_text = tmp_text.replace(gv.fieldsTitles[i], valuesFromCD[i])   
                            par.text=tmp_text
                        break

            if "Course DURATION" in par.text: 
                isCourseDuration = True

            #Remove comments start with "++" before Course Duration part
            if "++" in par.text and isCourseDuration == False : 
                if "Delete" in par.text: 
                    strIndex = par.text.find("++")
                    par.text = par.text.replace(par.text[strIndex:], "")
                else:
                    delete_paragraph(par)
            
            if self.rbCoursePerson.get() == "Lecturer":
                if "Course Coordinator" in par.text and isDeleted == False:
                    for i in range(6):
                        delete_paragraph(values[index+i])
                    isDeleted = True
            
            if self.rbCoursePerson.get() == "Course Coordinator":
                if "+Lecturer Name+" in par.text and isDeleted == False:
                    for i in range(6):
                        delete_paragraph(values[index+i])
                    isDeleted = True
            
            def delete_paragraph(paragraph):
                p = paragraph._element
                if p.getparent() != None: 
                    p.getparent().remove(p)
                    p._p = p._element = None 

#</editor-fold>

if __name__ == '__main__':
    root = Root()
    root.iconbitmap('img/icon.ico')
    logoImg = (Image.open("img/logo.jpg"))
    resizedImg = logoImg.resize((120,50), Image.ANTIALIAS)
    logoImg = ImageTk.PhotoImage(resizedImg)
    label = tk.Label(root, image=logoImg).place(x=280, y=10)
    root.bind('<Return>', root.funcEnter)

   #Built of UI:
    root.createHeader()
    root.createPeriod()
    root.createLecturerInf()
    root.endForm()

    root.mainloop()
